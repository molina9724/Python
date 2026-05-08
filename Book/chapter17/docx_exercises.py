# =====================================================================
#                    SECTION 3: WORD DOCUMENTS - BASICS
# =====================================================================

import os
import re
import sys
from datetime import datetime

import docx
import requests

# ----------------------------------------------------------------------
# 🟢 13: CREATE A NEW WORD DOCUMENT
#
# Learn: Document(), save()
#
# Tasks:
# 1. Create a new Document object
# 2. Save it with a filename
# 3. Verify the file was created
# 4. Open it manually to confirm it works
# ----------------------------------------------------------------------

PATH_FILE = (
    "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/doc_exercises.docx"
)

doc = docx.Document()
doc.save(PATH_FILE)

# ----------------------------------------------------------------------
# 🟢 14: ADD PARAGRAPHS
#
# Learn: add_paragraph()
#
# Tasks:
# 1. Create a new document
# 2. Add a paragraph with some text
# 3. Add two more paragraphs
# 4. Save the document
# 5. Open and verify content
# ----------------------------------------------------------------------

doc.add_paragraph("Paragraph 1")
doc.add_paragraph("Paragraph 2")
doc.add_paragraph("Paragraph 3")

doc.save(PATH_FILE)

# ----------------------------------------------------------------------
# 🟢 15: ADD HEADINGS
#
# Learn: add_heading()
#
# Tasks:
# 1. Create a new document
# 2. Add a main heading (level 0 or 1)
# 3. Add a subheading (level 2)
# 4. Add a paragraph under each heading
# 5. Save and verify the structure
# ----------------------------------------------------------------------

doc.add_heading("heading lvl0", 0)
doc.add_paragraph("paragraph under heading lvl10")

doc.add_heading("subheading", 2)
doc.add_paragraph("paragraph under heading lvl2")

doc.save(PATH_FILE)

# ----------------------------------------------------------------------
# 🟢 16: READ AN EXISTING DOCUMENT
#
# Learn: Document(filename), paragraphs
#
# Tasks:
# 1. Open an existing Word document
# 2. Access the paragraphs attribute
# 3. Print the number of paragraphs
# 4. Loop through and print each paragraph's text
# 5. Find a specific paragraph by content
# ----------------------------------------------------------------------

word = "under"
paragraphs_with_word = list()

print(f"The doc has {len(doc.paragraphs)} paragraphs")
for index, paragraph in enumerate(doc.paragraphs):
    print(paragraph.text)
    for run in paragraph.runs:
        if word in run.text:
            paragraphs_with_word.append(index + 1)

print(paragraphs_with_word)

# ----------------------------------------------------------------------
# 🟡 17: ACCESS PARAGRAPH PROPERTIES
#
# Learn: paragraph.text, paragraph.style
#
# Tasks:
# 1. Open a Word document
# 2. Get the first paragraph
# 3. Print its text content
# 4. Print its style name
# 5. Check if it's a heading style
# ----------------------------------------------------------------------

first_paragraph = doc.paragraphs[0]
print(first_paragraph.text)
print(first_paragraph.style)


# =====================================================================
#                    SECTION 4: WORD FORMATTING
# =====================================================================


# ----------------------------------------------------------------------
# 🟡 18: WORK WITH RUNS
#
# Learn: paragraph.runs, run.text
#
# Tasks:
# 1. Open or create a Word document
# 2. Access a paragraph
# 3. Get all runs in that paragraph
# 4. Print each run's text separately
# 5. Count the number of runs
# ----------------------------------------------------------------------

for run in first_paragraph.runs:
    print(run.text)

# ----------------------------------------------------------------------
# 🟡 19: FORMAT TEXT WITH RUNS
#
# Learn: run.bold, run.italic, run.underline
#
# Tasks:
# 1. Create a new document
# 2. Add a paragraph
# 3. Add a run with bold text
# 4. Add a run with italic text
# 5. Add a run with underlined text
# 6. Save and verify formatting
# ----------------------------------------------------------------------

paragraph_with_style = doc.add_paragraph("This paragraph is really hot\n")
paragraph_with_style.add_run("This is bold\n").bold = True
paragraph_with_style.add_run("This is italic\n").italic = True
paragraph_with_style.add_run("And this is underline, my man\n").underline = True

doc.save(PATH_FILE)

# ----------------------------------------------------------------------
# 🟡 20: ADD FORMATTED PARAGRAPH
#
# Learn: add_run() method
#
# Tasks:
# 1. Create a document
# 2. Add a paragraph with mixed formatting:
#    - Normal text, then bold, then italic
# 3. Use add_run() to add each formatted section
# 4. Save and check the result
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🟡 21: CHANGE FONT PROPERTIES
#
# Learn: run.font.size, run.font.name, run.font.color
#
# Tasks:
# 1. Create a document with a paragraph
# 2. Add a run
# 3. Change the font name (e.g., "Arial")
# 4. Change the font size (use Pt() from docx.shared)
# 5. Change the font color (use RGBColor from docx.shared)
# 6. Save and verify
# ----------------------------------------------------------------------

from docx.shared import Pt, RGBColor

for run in paragraph_with_style.runs:
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(r=0, g=200, b=50)

doc.save(PATH_FILE)

# ----------------------------------------------------------------------
# 🟡 22: APPLY PARAGRAPH STYLES
#
# Learn: Built-in styles
#
# Tasks:
# 1. Create a document
# 2. Add a paragraph with style='Title'
# 3. Add a paragraph with style='Heading 1'
# 4. Add a paragraph with style='Normal'
# 5. Add a paragraph with style='Quote'
# 6. Explore other available styles
# ----------------------------------------------------------------------

doc.add_paragraph("Title", "Title")
doc.add_paragraph("Heading 1", "Heading 1")
doc.add_paragraph("Normal", "Normal")
doc.add_paragraph("Quote", "Quote")
doc.add_paragraph("Caption", "Caption")

doc.save(PATH_FILE)

# =====================================================================
#                    SECTION 5: WORD DOCUMENT STRUCTURE
# =====================================================================


# ----------------------------------------------------------------------
# 🟡 23: ADD A TABLE
#
# Learn: add_table(), rows, columns, cell()
#
# Tasks:
# 1. Create a document
# 2. Add a table with 3 rows and 4 columns
# 3. Access cells and add text to them
# 4. Add header row content
# 5. Add data to remaining rows
# 6. Save and verify table structure
# ----------------------------------------------------------------------

table = doc.add_table(rows=4, cols=4)

header = table.rows[0].cells

values = ["Name", "Age", "Phone", "Blood"]
data = [
    ["CJ", 25, "123", "A+"],
    ["Andres", 50, "456", "AB+"],
    ["Carel", 12, "789", "O-"],
]

for index, cell in enumerate(header):
    cell.text = values[index]

for index, row in enumerate(table.rows):
    if index != 0:
        for index_2, cell in enumerate(row.cells):
            cell.text = str(data[index - 1][index_2])

doc.save(PATH_FILE)

# ----------------------------------------------------------------------
# 🟡 24: MODIFY TABLE CONTENT
#
# Learn: table.rows, table.columns, cell.text
#
# Tasks:
# 1. Open a document with a table
# 2. Access the first table
# 3. Loop through all rows and print content
# 4. Modify a specific cell's content
# 5. Add a new row to the table
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🟡 25: ADD IMAGES
#
# Learn: add_picture()
#
# Tasks:
# 1. Create a document
# 2. Add a heading
# 3. Add a paragraph
# 4. Add an image from a file
# 5. Set image width (use Inches() from docx.shared)
# 6. Save and verify image appears
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🟡 26: ADD PAGE BREAKS
#
# Learn: add_page_break()
#
# Tasks:
# 1. Create a document
# 2. Add some content to page 1
# 3. Add a page break
# 4. Add content that should appear on page 2
# 5. Add another page break and more content
# 6. Save and verify page breaks work
# ----------------------------------------------------------------------
