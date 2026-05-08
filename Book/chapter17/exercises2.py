import docx
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Inches

BASE_PATH = "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/"
FILE_PATH = "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/demo.docx"
doc = docx.Document(FILE_PATH)

for paragraph in doc.paragraphs:
    print(paragraph.text)
    for run in paragraph.runs:
        print(run.text)

full_text = []
for paragraph in doc.paragraphs:
    full_text.append(paragraph.text)
full_string_text = "\n\n".join(full_text)

print(full_text)
print(full_string_text)

print(doc.paragraphs[0].style)
doc.paragraphs[0].style = "Normal"

for run in doc.paragraphs[0].runs:
    run.font.strike = True
    run.font.imprint = True
    run.font.outline = True
    run.font.shadow = True

doc.save("/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/restyled.docx")

my_doc = docx.Document()
my_doc.add_paragraph("Say hello to the hackers")

paragraph_obj_2 = my_doc.add_paragraph("This is the second paragraph")
paragraph_obj_3 = my_doc.add_paragraph("This is the third paragraph")

paragraph_obj_2.add_run("\nCheck how I add at the end of the second paragraph")
paragraph_obj_3.add_run("Just testing")

my_doc.add_paragraph("I'm the ruler here", "Title")

my_doc.add_heading("Header 1", 1)
my_doc.add_heading("Header 2", 2)
my_doc.add_heading("Header 3", 3)
my_doc.add_heading("Header 4", 4)

# These bad boys always return the object they added so you don't have to extract them

print(my_doc.paragraphs[7].runs[0].text)

my_doc.paragraphs[7].runs[0].add_break(WD_BREAK.PAGE)
my_doc.add_paragraph("This is the second page we don't fuck around")

print(len(my_doc.paragraphs))

my_doc.add_picture(
    "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/zophie.png",
    width=Inches(1),
    height=Cm(4),
)

my_doc.save(BASE_PATH + "my_doc.docx")
