# ======================================================================
# 📄 PDF AND WORD DOCUMENTS EXERCISES - Document Automation
# ======================================================================
# Practice exercises - Write everything from scratch!
# Prerequisites: PyPDF2 and python-docx installed
# ======================================================================

# pip install PyPDF2 python-docx

import pypdf
from docx import Document

# =====================================================================
#                    SECTION 1: PDF BASICS - READING
# =====================================================================


# ----------------------------------------------------------------------
# 🟢 1: OPEN AND READ A PDF
#
# Learn: PdfReader(), len(), pages
#
# Tasks:
# 1. Open a PDF file in read-binary mode ('rb')
# 2. Create a PdfReader object
# 3. Print how many pages the PDF has
# 4. Close the file properly
# ----------------------------------------------------------------------

BASE_PATH = "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/"

PDF_FILE = (
    "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/Recursion_Chapter1.pdf"
)

print("Hello")

with open(
    PDF_FILE,
    "rb",
) as file:
    print("File open")
print("File closed")

reader = pypdf.PdfReader(PDF_FILE)
print(len(reader.pages))
# reader.close()

# ----------------------------------------------------------------------
# 🟢 2: EXTRACT TEXT FROM A PAGE
#
# Learn: pages[], extract_text()
#
# Tasks:
# 1. Open a PDF file
# 2. Access the first page (index 0)
# 3. Extract text from that page
# 4. Print the extracted text
# 5. Extract text from page 2 (if it exists)
# ----------------------------------------------------------------------

first_page = reader.pages[0]
first_page_text = first_page.extract_text()
print(first_page_text)

try:
    second_page = reader.pages[1]
    second_page_text = second_page.extract_text()
    print(second_page_text)
except Exception as e:
    print(f"There was an error: {e}")

# ----------------------------------------------------------------------
# 🟢 3: EXTRACT TEXT FROM ALL PAGES
#
# Learn: Iterating through PDF pages
#
# Tasks:
# 1. Open a multi-page PDF
# 2. Loop through all pages
# 3. Extract and print text from each page
# 4. Store all text in a single string variable
# 5. Count total characters extracted
# ----------------------------------------------------------------------

all_text = ""
for index, page in enumerate(reader.pages, 1):
    print(f"Page {index}, {page.extract_text()}")
    all_text += page.extract_text()
print(len(all_text))


# ----------------------------------------------------------------------
# 🟡 4: CHECK PDF PROPERTIES
#
# Learn: metadata attribute
#
# Tasks:
# 1. Open a PDF file
# 2. Access the PDF metadata
# 3. Print the author (if available)
# 4. Print the title (if available)
# 5. Print all available metadata keys
# ----------------------------------------------------------------------

metadata = reader.metadata

if metadata is not None:
    try:
        print(metadata["/author"])
        print(metadata["/title"])
    except Exception as e:
        print(f"There was an error: {e}")

    for key, value in metadata.items():
        print(f"{key} : {value}")

# ----------------------------------------------------------------------
# 🟡 5: CHECK IF PDF IS ENCRYPTED
#
# Learn: is_encrypted property
#
# Tasks:
# 1. Open a PDF file
# 2. Check if the PDF is encrypted
# 3. Print appropriate message based on encryption status
# 4. If encrypted, attempt to decrypt with a password
# 5. Handle the case where decryption fails
# ----------------------------------------------------------------------

if reader.is_encrypted:
    print("The PDF is encrypted")
else:
    print("The PDF is not encrypted")

ENCRYPTED_PDF_FILE = (
    "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/with_watermark.pdf"
)
encrypted_reader = pypdf.PdfReader(ENCRYPTED_PDF_FILE)

try:
    encrypted_reader.decrypt("wrong_password")
except Exception as e:
    print(f"There was an error unencrypting the file: {e}")
    print("But now it won't :)")
    encrypted_reader.decrypt("swordfish")
    print(f"The encrypted file has {len(encrypted_reader.pages)} pages")

# =====================================================================
#                    SECTION 2: PDF MANIPULATION
# =====================================================================


# ----------------------------------------------------------------------
# 🟢 6: CREATE A NEW PDF (COPY PAGES)
#
# Learn: PdfWriter(), add_page(), write()
#
# Tasks:
# 1. Open an existing PDF for reading
# 2. Create a PdfWriter object
# 3. Copy the first page to the writer
# 4. Save as a new PDF file
# 5. Verify the new file was created
# ----------------------------------------------------------------------

writer = pypdf.PdfWriter()
writer.append(PDF_FILE, [0])

with open(BASE_PATH + "exercise6.pdf", "wb") as file:
    writer.write(file)

EXERCISE6 = "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/exercise6.pdf"

writer = pypdf.PdfWriter()
writer.append(PDF_FILE, [1])
with open(BASE_PATH + "exercise6a.pdf", "wb") as file:
    writer.write(file)

EXERCISE6a = (
    "/Users/daniel_molina/Downloads/Python/Python/Book/chapter17/exercise6a.pdf"
)


# ----------------------------------------------------------------------
# 🟡 7: MERGE MULTIPLE PDFS
#
# Learn: PdfMerger(), append()
#
# Tasks:
# 1. Create a PdfMerger object
# 2. Append two or more PDF files
# 3. Write the merged result to a new file
# 4. Verify the page count of merged PDF
# 5. Close the merger properly
# ----------------------------------------------------------------------

writer = pypdf.PdfWriter()
writer.append(EXERCISE6)
writer.append(EXERCISE6a)

with open(BASE_PATH + "exercise7.pdf", "wb") as file:
    writer.write(file)

# ----------------------------------------------------------------------
# 🟡 8: EXTRACT SPECIFIC PAGES
#
# Learn: Selecting pages to copy
#
# Tasks:
# 1. Open a multi-page PDF
# 2. Create a PdfWriter object
# 3. Copy only pages 2 and 4 (or specific pages you choose)
# 4. Save as a new PDF
# 5. Verify the new PDF has correct number of pages
# ----------------------------------------------------------------------

writer = pypdf.PdfWriter()
EXERCISE8 = BASE_PATH + "exercise8.pdf"

writer.append(PDF_FILE, [2, 4])
with open(EXERCISE8, "wb") as file:
    writer.write(file)

assert len(pypdf.PdfReader(EXERCISE8).pages) == 2

# ----------------------------------------------------------------------
# 🟡 9: ROTATE PAGES
#
# Learn: rotate(90), rotate(180), rotate(270)
#
# Tasks:
# 1. Open a PDF file
# 2. Get a specific page
# 3. Rotate the page 90 degrees clockwise
# 4. Create a writer and add the rotated page
# 5. Save and verify rotation worked
# ----------------------------------------------------------------------

writer = pypdf.PdfWriter()
writer.append(PDF_FILE, [10])
writer.pages[0].rotate(90)

with open(BASE_PATH + "exercise10.pdf", "wb") as file:
    writer.write(file)

# ----------------------------------------------------------------------
# 🟡 10: SPLIT A PDF
#
# Learn: Creating multiple PDFs from one
#
# Tasks:
# 1. Open a PDF with multiple pages
# 2. Create separate PDFs for each page
# 3. Name each file with page number (page_1.pdf, page_2.pdf, etc.)
# 4. Verify all files were created
# 5. Count total files created
# ----------------------------------------------------------------------

writer = pypdf.PdfWriter()

total_pages = len(pypdf.PdfReader(PDF_FILE).pages)
for i in range(total_pages):
    with open(BASE_PATH + f"/exercise10/page_{i+1}.pdf", "wb") as f:
        aux_writer = pypdf.PdfWriter()
        aux_writer.append(PDF_FILE, [i])
        aux_writer.write(f)

# ----------------------------------------------------------------------
# 🟡 11: ENCRYPT A PDF
#
# Learn: encrypt() method
#
# Tasks:
# 1. Open a PDF file
# 2. Create a PdfWriter and copy all pages
# 3. Encrypt with a user password
# 4. Optionally add an owner password
# 5. Save and test that password is required
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🟡 12: DECRYPT A PDF
#
# Learn: decrypt() method
#
# Tasks:
# 1. Open an encrypted PDF
# 2. Check if it's encrypted
# 3. Decrypt using the correct password
# 4. Extract and print text after decryption
# 5. Handle incorrect password gracefully
# ----------------------------------------------------------------------


# =====================================================================
#                    SECTION 6: REAL-WORLD SCENARIOS
# =====================================================================


# ----------------------------------------------------------------------
# 🔴 27: PDF TEXT SEARCH
#
# Scenario: Search for specific text across multiple PDFs
#
# Tasks:
# 1. Get a list of all PDF files in a directory
# 2. Write a function to search for a keyword in a PDF
# 3. Search all PDFs for a specific term
# 4. Print which PDFs contain the term and on which page
# 5. Save results to a text file
# 6. Handle errors for corrupted/unreadable PDFs
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 28: PDF COMBINER TOOL
#
# Scenario: Combine selected PDFs in a specific order
#
# Tasks:
# 1. Create a list of PDF filenames to combine
# 2. Verify all files exist before starting
# 3. Merge PDFs in the specified order
# 4. Add a cover page (first page of combined doc)
# 5. Save with a meaningful filename including date
# 6. Print summary of pages combined
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 29: INVOICE GENERATOR
#
# Scenario: Generate Word document invoices
#
# Tasks:
# 1. Create a function that takes invoice data as parameters
# 2. Generate a document with company header
# 3. Add invoice number, date, and customer info
# 4. Create a table with items, quantities, and prices
# 5. Calculate and display total
# 6. Save with invoice number in filename
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 30: REPORT TEMPLATE FILLER
#
# Scenario: Fill in a report template with data
#
# Tasks:
# 1. Create a Word template with placeholder text
#    (e.g., {{NAME}}, {{DATE}}, {{CONTENT}})
# 2. Write a function to open and read the template
# 3. Replace placeholders with actual data
# 4. Preserve formatting of the original template
# 5. Save as a new document (keep template intact)
# 6. Generate multiple reports from a data list
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 31: PDF WATERMARKER
#
# Scenario: Add a watermark to all pages of a PDF
#
# Tasks:
# 1. Create or obtain a watermark PDF (single page)
# 2. Open the source PDF
# 3. Merge watermark onto each page
# 4. Save as a new watermarked PDF
# 5. Handle PDFs of different page sizes
# 6. Make watermark opacity adjustable (if possible)
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 32: DOCUMENT CONVERTER
#
# Scenario: Batch process documents
#
# Tasks:
# 1. Get all Word documents in a folder
# 2. Extract text from each document
# 3. Create a summary document with all filenames and word counts
# 4. Find the longest and shortest documents
# 5. Create a combined "all text" document
# 6. Generate a report of the batch processing
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 33: MEETING MINUTES GENERATOR
#
# Scenario: Create formatted meeting minutes
#
# Tasks:
# 1. Create a function taking meeting details as parameters
# 2. Add formatted header with date, time, attendees
# 3. Add agenda items as headings
# 4. Add discussion points under each agenda item
# 5. Add action items table (who, what, when)
# 6. Add signature lines at the bottom
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 34: PDF PAGE EXTRACTOR GUI
#
# Scenario: Extract specific page ranges from a PDF
#
# Tasks:
# 1. Write function accepting filename and page range string
#    (e.g., "1-3, 5, 7-9")
# 2. Parse the page range string
# 3. Extract specified pages
# 4. Handle invalid page numbers gracefully
# 5. Save extracted pages to new PDF
# 6. Print summary of pages extracted
# ----------------------------------------------------------------------


# ----------------------------------------------------------------------
# 🔴 35: BULK DOCUMENT CREATOR
#
# Scenario: Create personalized documents from a template
#
# Tasks:
# 1. Create a CSV file with recipient data (name, address, etc.)
# 2. Create a Word template with placeholders
# 3. Read the CSV file
# 4. For each row, create a personalized document
# 5. Save each with recipient's name in filename
# 6. Create a log of all documents generated
# ----------------------------------------------------------------------


# ======================================================================
# 📄 QUICK REFERENCE - PyPDF2 Reading
# ======================================================================
#
# Opening PDFs:
#   pdf_file = open('file.pdf', 'rb')
#   reader = PyPDF2.PdfReader(pdf_file)
#
# Properties:
#   len(reader.pages)         # Number of pages
#   reader.metadata           # PDF metadata dict
#   reader.is_encrypted       # Check if encrypted
#
# Reading Pages:
#   page = reader.pages[0]    # Get first page (0-indexed)
#   text = page.extract_text() # Extract text from page
#
# Decryption:
#   reader.decrypt('password')  # Decrypt with password
#
# ======================================================================


# ======================================================================
# 📄 QUICK REFERENCE - PyPDF2 Writing
# ======================================================================
#
# Creating PDFs:
#   writer = PyPDF2.PdfWriter()
#
# Adding Pages:
#   writer.add_page(page)     # Add a page object
#
# Rotating Pages:
#   page.rotate(90)           # Rotate clockwise
#   page.rotate(180)          # Rotate 180 degrees
#   page.rotate(270)          # Rotate counter-clockwise
#
# Encryption:
#   writer.encrypt('user_pwd', 'owner_pwd')
#
# Saving:
#   with open('output.pdf', 'wb') as output:
#       writer.write(output)
#
# ======================================================================


# ======================================================================
# 📄 QUICK REFERENCE - PyPDF2 Merging
# ======================================================================
#
# Merging PDFs:
#   merger = PyPDF2.PdfMerger()
#   merger.append('file1.pdf')
#   merger.append('file2.pdf')
#   merger.write('merged.pdf')
#   merger.close()
#
# Merge specific pages:
#   merger.append('file.pdf', pages=(0, 3))  # Pages 1-3
#
# ======================================================================


# ======================================================================
# 📄 QUICK REFERENCE - python-docx Basics
# ======================================================================
#
# Creating/Opening:
#   doc = Document()              # New document
#   doc = Document('file.docx')   # Open existing
#
# Adding Content:
#   doc.add_heading('Title', 0)   # Main title
#   doc.add_heading('Heading', 1) # Heading level 1
#   doc.add_paragraph('Text')     # Normal paragraph
#   doc.add_page_break()          # Page break
#
# Saving:
#   doc.save('filename.docx')
#
# ======================================================================


# ======================================================================
# 📄 QUICK REFERENCE - python-docx Reading
# ======================================================================
#
# Paragraphs:
#   doc.paragraphs                # List of all paragraphs
#   para.text                     # Paragraph text content
#   para.style.name               # Paragraph style name
#
# Runs (text segments):
#   para.runs                     # List of runs in paragraph
#   run.text                      # Run text content
#   run.bold                      # True/False/None
#   run.italic                    # True/False/None
#   run.underline                 # True/False/None
#
# ======================================================================


# ======================================================================
# 📄 QUICK REFERENCE - python-docx Formatting
# ======================================================================
#
# Adding formatted text:
#   para = doc.add_paragraph()
#   run = para.add_run('Bold text')
#   run.bold = True
#
# Font properties:
#   from docx.shared import Pt, Inches, RGBColor
#
#   run.font.name = 'Arial'
#   run.font.size = Pt(12)
#   run.font.color.rgb = RGBColor(255, 0, 0)  # Red
#
# Paragraph with style:
#   doc.add_paragraph('Text', style='Quote')
#
# ======================================================================


# ======================================================================
# 📄 QUICK REFERENCE - python-docx Tables & Images
# ======================================================================
#
# Tables:
#   table = doc.add_table(rows=3, cols=4)
#   cell = table.cell(0, 0)       # Row 0, Column 0
#   cell.text = 'Content'
#   row = table.rows[0]           # First row
#   table.add_row()               # Add new row
#
# Images:
#   doc.add_picture('image.png')
#   doc.add_picture('image.png', width=Inches(2))
#   doc.add_picture('image.png', height=Inches(3))
#
# ======================================================================


# ======================================================================
# 🚀 SETUP INSTRUCTIONS
# ======================================================================
#
# Installation:
#   pip install PyPDF2
#   pip install python-docx
#
# Imports:
#   import PyPDF2
#   from docx import Document
#   from docx.shared import Pt, Inches, RGBColor
#
# Sample PDFs for practice:
#   - Create your own using Word -> Save as PDF
#   - Use any PDF files you have on your computer
#   - Download free sample PDFs online
#
# Note on PDF text extraction:
#   - Works best on text-based PDFs
#   - Scanned PDFs (images) won't extract text
#   - Formatting may not be perfectly preserved
#
# ======================================================================
